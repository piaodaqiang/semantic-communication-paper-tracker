from __future__ import annotations

import html
import zipfile
from collections import Counter
from pathlib import Path

from .schema import PAPER_FIELDS, Paper


def export_excel(path: Path, papers: list[Paper]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ModuleNotFoundError:
        fallback = path.with_suffix(".csv")
        from .storage import save_csv

        save_csv(fallback, papers)
        return

    wb = Workbook()
    ws = wb.active
    ws.title = "papers"
    ws.append(PAPER_FIELDS)
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = header_fill
    for paper in papers:
        row = paper.to_dict()
        ws.append([row.get(field, "") for field in PAPER_FIELDS])
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for index, field in enumerate(PAPER_FIELDS, start=1):
        width = 14
        if field in {"title", "abstract", "paper_url", "pdf_url", "code_url", "project_url"}:
            width = 42
        elif field in {"authors", "application_scenario", "technical_framework"}:
            width = 24
        ws.column_dimensions[get_column_letter(index)].width = width
    wb.save(path)


def export_markdown_summary(path: Path, papers: list[Paper], title: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    lines = [f"# {title}", ""]
    lines.extend(build_summary_lines(papers))
    lines.append("")
    lines.append("## 建议精读论文")
    for paper in sorted(papers, key=lambda item: item.relevance_score, reverse=True)[:10]:
        code = "开源" if paper.is_open_source else "未发现开源"
        lines.append(f"- {paper.title} ({paper.year}) | {paper.application_scenario} | {paper.technical_framework} | {code}")
    lines.append("")
    lines.append("## 逐篇简表")
    lines.append("| 序号 | 题名 | 年份 | 场景 | 技术框架 | 相关性 | 开源 |")
    lines.append("| --- | --- | --- | --- | --- | --- | --- |")
    for index, paper in enumerate(sorted(papers, key=lambda item: item.relevance_score, reverse=True), start=1):
        code = paper.code_url if paper.is_open_source else "未发现"
        lines.append(
            f"| {index} | {paper.title} | {paper.year} | {paper.application_scenario} | "
            f"{paper.technical_framework} | {paper.relevance_score} | {code} |"
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def export_word_report(path: Path, papers: list[Paper], title: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document
    except ModuleNotFoundError:
        write_minimal_docx(path, papers, title)
        return

    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_heading("一、总体概况", level=1)
    for line in build_summary_lines(papers):
        doc.add_paragraph(line)
    doc.add_paragraph("本周报告仅纳入满足语义通信核心关键词门槛的 curated 论文；低相关候选保留在 Daily Inbox 中，不进入导师汇报文件。")

    doc.add_heading("二、分类统计", level=1)
    add_counter_table(doc, "应用场景", Counter(paper.application_scenario for paper in papers))
    add_counter_table(doc, "技术框架", Counter(paper.technical_framework for paper in papers))

    doc.add_heading("建议精读论文", level=1)
    for paper in sorted(papers, key=lambda item: item.relevance_score, reverse=True)[:10]:
        doc.add_paragraph(
            f"{paper.title}（{paper.year}）｜{paper.application_scenario}｜{paper.technical_framework}｜相关性 {paper.relevance_score}",
            style="List Bullet",
        )
    doc.add_heading("三、逐篇明细", level=1)
    add_paper_table(doc, papers)

    doc.add_heading("四、开源情况", level=1)
    open_papers = [item for item in papers if item.is_open_source]
    if open_papers:
        for paper in open_papers[:20]:
            doc.add_paragraph(f"{paper.title}：{paper.code_url}", style="List Bullet")
    else:
        doc.add_paragraph("本次 curated 论文中暂未检测到明确代码仓库链接，后续可人工补充项目主页或 GitHub 仓库。")
    doc.save(path)


def build_summary_lines(papers: list[Paper]) -> list[str]:
    year_counter = Counter(str(paper.year) for paper in papers if paper.year)
    scenario_counter = Counter(paper.application_scenario for paper in papers)
    framework_counter = Counter(paper.technical_framework for paper in papers)
    open_count = sum(1 for paper in papers if paper.is_open_source)
    return [
        f"论文总数：{len(papers)}",
        f"开源论文数：{open_count}",
        "年份分布：" + format_counter(year_counter),
        "场景分布：" + format_counter(scenario_counter),
        "技术框架分布：" + format_counter(framework_counter),
    ]


def format_counter(counter: Counter[str]) -> str:
    if not counter:
        return "无"
    return "；".join(f"{key} {value}" for key, value in counter.most_common())


def write_minimal_docx(path: Path, papers: list[Paper], title: str) -> None:
    detail_rows = [["序号", "题名", "年份", "应用场景", "技术框架", "相关性", "开源"]]
    for index, paper in enumerate(sorted(papers, key=lambda item: item.relevance_score, reverse=True), start=1):
        detail_rows.append(
            [
                str(index),
                paper.title,
                str(paper.year or ""),
                paper.application_scenario,
                paper.technical_framework,
                str(paper.relevance_score),
                paper.code_url if paper.is_open_source else "未发现",
            ]
        )
    scenario_rows = [["应用场景", "数量"], *[[key, str(value)] for key, value in Counter(p.application_scenario for p in papers).most_common()]]
    framework_rows = [["技术框架", "数量"], *[[key, str(value)] for key, value in Counter(p.technical_framework for p in papers).most_common()]]
    recommended = [
        f"{paper.title}（{paper.year}）｜{paper.application_scenario}｜{paper.technical_framework}｜相关性 {paper.relevance_score}"
        for paper in sorted(papers, key=lambda item: item.relevance_score, reverse=True)[:10]
    ]
    open_source_lines = (
        [f"{paper.title}：{paper.code_url}" for paper in [item for item in papers if item.is_open_source][:20]]
        or ["本次 curated 论文中暂未检测到明确代码仓库链接，后续可人工补充项目主页或 GitHub 仓库。"]
    )
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>'
        + _docx_block(title)
        + _docx_block("一、总体概况")
        + "".join(_docx_block(text) for text in build_summary_lines(papers))
        + _docx_block("本周报告仅纳入满足语义通信核心关键词门槛的 curated 论文；低相关候选保留在 Daily Inbox 中，不进入导师汇报文件。")
        + _docx_block("二、分类统计")
        + _docx_table(scenario_rows)
        + _docx_table(framework_rows)
        + _docx_block("三、建议精读论文")
        + "".join(_docx_block(text) for text in recommended)
        + _docx_block("四、逐篇明细")
        + _docx_table(detail_rows)
        + _docx_block("五、开源情况")
        + "".join(_docx_block(text) for text in open_source_lines)
        + "<w:sectPr/></w:body></w:document>"
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", CONTENT_TYPES)
        docx.writestr("_rels/.rels", RELS)
        docx.writestr("word/document.xml", document_xml)


def _docx_paragraph(text: str) -> str:
    return f"<w:p><w:r><w:t>{html.escape(text)}</w:t></w:r></w:p>"


def _docx_block(text: str) -> str:
    style = ""
    if text and text[0] in {"一", "二", "三", "四", "五"} and "、" in text[:3]:
        style = '<w:pPr><w:pStyle w:val="Heading1"/></w:pPr>'
    return f"<w:p>{style}<w:r><w:t>{html.escape(text)}</w:t></w:r></w:p>"


def _docx_table(rows: list[list[str]]) -> str:
    grid = "<w:tblGrid>" + "".join('<w:gridCol w:w="2400"/>' for _ in rows[0]) + "</w:tblGrid>"
    body = []
    for row in rows:
        cells = []
        for value in row:
            cells.append(
                "<w:tc><w:tcPr><w:tcW w:w=\"2400\" w:type=\"dxa\"/></w:tcPr>"
                f"{_docx_paragraph(str(value))}</w:tc>"
            )
        body.append("<w:tr>" + "".join(cells) + "</w:tr>")
    return (
        '<w:tbl><w:tblPr><w:tblStyle w:val="TableGrid"/><w:tblW w:w="0" w:type="auto"/>'
        '<w:tblBorders><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/>'
        '<w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/>'
        '<w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders>'
        f"</w:tblPr>{grid}{''.join(body)}</w:tbl>"
    )


def add_counter_table(doc, label: str, counter: Counter[str]) -> None:
    doc.add_paragraph(label)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "类别"
    table.rows[0].cells[1].text = "数量"
    for key, value in counter.most_common():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)


def add_paper_table(doc, papers: list[Paper]) -> None:
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["序号", "题名", "年份", "应用场景", "技术框架", "相关性", "开源"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for index, paper in enumerate(sorted(papers, key=lambda item: item.relevance_score, reverse=True), start=1):
        row = table.add_row().cells
        values = [
            str(index),
            paper.title,
            str(paper.year or ""),
            paper.application_scenario,
            paper.technical_framework,
            str(paper.relevance_score),
            paper.code_url if paper.is_open_source else "未发现",
        ]
        for cell_index, value in enumerate(values):
            row[cell_index].text = value


CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
